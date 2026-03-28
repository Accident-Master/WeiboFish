import io
import re

import matplotlib.font_manager as fm
import matplotlib.pyplot as plt
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

def extract_id(id_str):
    if id_str is None or str(id_str).lower() == 'null': return None
    nums = re.findall(r'\d+', str(id_str))
    return int(nums[0]) if nums else None


def set_matplotlib_font():
    fonts = ['Microsoft YaHei', 'SimHei', 'Arial Unicode MS']
    available = [f.name for f in fm.fontManager.ttflist]
    for f in fonts:
        if f in available:
            plt.rcParams['font.sans-serif'] = [f]
            break
    plt.rcParams['axes.unicode_minus'] = False
    plt.rcParams['figure.dpi'] = 600
    plt.rcParams['savefig.dpi'] = 600


def create_word_report(city_name, report_text):
    doc = Document()
    doc.styles['Normal'].font.name = u'微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    doc.styles['Normal'].font.size = Pt(11)

    title = doc.add_heading(f'【决策内参】针对{city_name}舆论演化动力学之专报', level=1)
    title.alignment = 1
    for run in title.runs:
        run.font.name = u'微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    doc.add_paragraph()

    for line in report_text.split('\n'):
        line = line.strip()
        if not line:
            continue

        p = doc.add_paragraph()
        parts = line.split('**')
        for i, text in enumerate(parts):
            run = p.add_run(text)
            if i % 2 != 0:
                run.bold = True
                run.font.color.rgb = RGBColor(194, 53, 49)
            run.font.name = u'微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
