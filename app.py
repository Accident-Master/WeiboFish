import sys
import json
import random
import math
import re
import io
import pandas as pd
from pathlib import Path
import streamlit as st
import os
from datetime import datetime
os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"
import matplotlib.font_manager as font_manager
import matplotlib.pyplot as plt
font_manager.fontManager.addfont('MSYH.TTC')
plt.rcParams['font.sans-serif'] = ['Microsoft YaHei']
plt.rcParams['axes.unicode_minus'] = False
import matplotlib.font_manager as fm
import matplotlib.gridspec as gridspec
import networkx as nx
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

# ==========================================
# 0. API 配置与使用监控 (新增)
# ==========================================
MY_API_KEY = ""

def log_usage():
    log_file = "usage_log.csv"
    time_now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    new_entry = pd.DataFrame({"使用时间": [time_now], "调用状态": ["成功"]})
    if not os.path.exists(log_file):
        new_entry.to_csv(log_file, index=False, encoding='utf-8-sig')
    else:
        new_entry.to_csv(log_file, mode='a', header=False, index=False, encoding='utf-8-sig')

def get_total_usage():
    log_file = "usage_log.csv"
    if os.path.exists(log_file):
        try:
            df = pd.read_csv(log_file)
            return len(df)
        except:
            return 0
    return 0

# ==========================================
# 1. 路径与环境初始化
# ==========================================
current_dir = Path(__file__).resolve().parent
if str(current_dir) not in sys.path:
    sys.path.insert(0, str(current_dir))

from src.features.models.text_analyzer import WeiboFeatureExtractor
from src.config.load_params import ReactionModel
from src.sim.memory import HistoricalMemory

# ==========================================
# 2. 页面与 UI 配置
# ==========================================
st.set_page_config(page_title="weibofish 舆情沙盘", page_icon="🐟", layout="wide")

if "sim_completed" not in st.session_state:
    st.session_state.sim_completed = False

st.markdown("""
<style>
    /* 1. 最精准的手术刀：只隐藏 Deploy 按钮和三点菜单，绝对不碰任何顶层框架！ */
    .stAppDeployButton { display: none !important; }
    .stDeployButton { display: none !important; }
    #MainMenu { display: none !important; }
    
    /* 2. 隐藏底部水印 */
    footer { display: none !important; }

    /* 👇 下面是沙盘原有的核心业务样式，保持原封不动 👇 */
    .report-box {
        background-color: #f8f9fa; color: #333333; padding: 25px;
        border-radius: 8px; border-left: 6px solid #1c4e7d;
        font-family: 'Microsoft YaHei', 'SimSun', serif;
        line-height: 1.8; font-size: 16px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.05);
    }
    .report-title { color: #c23531; font-weight: bold; font-size: 22px; margin-bottom: 15px; }
    div[data-testid="metric-container"] > div > div > div { font-size: 1.8rem !important; color: #1c4e7d; }
    .os-box { background-color: #2b2b2b; color: #4af55d; padding: 10px; border-radius: 6px; margin-bottom: 10px; font-family: monospace; font-size: 13px; border-left: 4px solid #4af55d;}
    .os-view { border-left: 4px solid #888888; color: #aaaaaa; } 
    .weibo-comment-main { background-color: #ffffff; padding: 15px; border-radius: 8px; margin-bottom: 15px; border: 1px solid #e0e0e0; box-shadow: 0 1px 3px rgba(0,0,0,0.05); }
    .weibo-comment-sub { background-color: #f9f9f9; padding: 12px; margin-top: 10px; margin-left: 20px; border-radius: 6px; border-left: 3px solid #00b4d8; }
    .comment-header { font-weight: bold; color: #eb7350; font-size: 14px; margin-bottom: 5px; }
    .comment-traits { font-weight: normal; color: #999; font-size: 12px; margin-left: 8px; }
    .comment-content { color: #333; font-size: 15px; line-height: 1.5; }
    .comment-actions { color: #808080; font-size: 13px; margin-top: 8px; display: flex; align-items: center; gap: 15px; }
</style>
""", unsafe_allow_html=True)


# ==========================================
# 3. 核心数据与缓存加载器
# ==========================================
def read_csv_safe(file_path):
    encodings = ['utf-8', 'gbk', 'gb18030', 'utf-8-sig']
    for enc in encodings:
        try:
            return pd.read_csv(file_path, encoding=enc)
        except UnicodeDecodeError:
            continue
    return pd.read_csv(file_path, encoding='utf-8', errors='ignore')


@st.cache_data(show_spinner=False)
def load_agenda_data():
    data_dir = current_dir / "data"
    try:
        mapping_files = list(data_dir.glob("*账号匹配*.csv"))
        vol_files = list(data_dir.glob("*议程波动_账号汇总*.csv"))

        if not mapping_files:
            raise FileNotFoundError("找不到 账号匹配 CSV 文件")

        df_map = read_csv_safe(mapping_files[0]).dropna(subset=['省份', '城市'])

        prov_city_dict = df_map.groupby('省份')['城市'].unique().apply(list).to_dict()
        city_vol_dict = {c: 0.2 for c in df_map['城市'].unique()}
        prov_vol_dict = {p: 0.2 for p in df_map['省份'].unique()}
        global_vol = 0.2

        # --- 尝试关联波动数据 ---
        if vol_files:
            df_vol = read_csv_safe(vol_files[0])
            if '账号名' in df_map.columns and 'account' in df_vol.columns:
                df = pd.merge(df_map, df_vol, left_on='账号名', right_on='account', how='inner')
                if not df.empty:
                    prov_city_dict = df.groupby('省份')['城市'].unique().apply(list).to_dict()
                    city_vol_dict = df.groupby('城市')['vol_tv_mean'].mean().to_dict()
                    prov_vol_dict = df.groupby('省份')['vol_tv_mean'].mean().to_dict()
                    global_vol = df['vol_tv_mean'].mean()
            else:
                pass

        return prov_city_dict, city_vol_dict, prov_vol_dict, global_vol

    except Exception as e:
        st.error(f"⚠️ 数据加载失败: {e}")
        # 最终保底方案
        return {"默认省份": ["默认城市"]}, {"默认城市": 0.2}, {"默认省份": 0.2}, 0.2

@st.cache_resource(show_spinner=False)
def load_ai_engines():
    nlp = WeiboFeatureExtractor()
    stats = ReactionModel()
    mem = HistoricalMemory()
    with open(current_dir / "data" / "agent_personas.json", "r", encoding="utf-8") as f:
        personas = json.load(f)
    return nlp, stats, mem, personas


def extract_id(id_str):
    if id_str is None or str(id_str).lower() == 'null': return None
    nums = re.findall(r'\d+', str(id_str))
    return int(nums[0]) if nums else None


def set_matplotlib_font():
    fonts = ['Microsoft YaHei', 'SimHei', 'Arial Unicode MS']
    available = [f.name for f in fm.fontManager.ttflist]
    for f in fonts:
        if f in available:
            plt.rcParams['font.sans-serif'] = [f]
            break
    plt.rcParams['axes.unicode_minus'] = False
    plt.rcParams['figure.dpi'] = 600
    plt.rcParams['savefig.dpi'] = 600


def create_word_report(city_name, report_text):
    doc = Document()
    doc.styles['Normal'].font.name = u'微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    doc.styles['Normal'].font.size = Pt(11)

    title = doc.add_heading(f'【决策内参】针对{city_name}舆论演化动力学之专报', level=1)
    title.alignment = 1
    for run in title.runs:
        run.font.name = u'微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    doc.add_paragraph()

    for line in report_text.split('\n'):
        line = line.strip()
        if not line:
            continue

        p = doc.add_paragraph()
        parts = line.split('**')
        for i, text in enumerate(parts):
            run = p.add_run(text)
            if i % 2 != 0:
                run.bold = True
                run.font.color.rgb = RGBColor(194, 53, 49)
            run.font.name = u'微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ==========================================
# 4. Streamlit 适配的 Agent 类 (两步互动独立解耦版)
# ==========================================
class StreamlitAgent:
    def __init__(self, agent_id, persona, client):
        self.agent_id = agent_id
        self.persona = persona
        self.trust = persona['initial_trust']
        self.role = persona['demographics']['occupation']
        self.is_exposed = False
        self.has_interacted = False
        self.client = client

    def react(self, post, history, social_context, bias_str, interact_post, interact_comment, location, is_local_mode):

        local_persona = dict(self.persona)
        if is_local_mode and 'demographics' in local_persona:
            local_demo = dict(local_persona['demographics'])
            local_demo['location'] = location
            local_persona['demographics'] = local_demo

        if not interact_post and not interact_comment:
            action_prompt = """【强制行为指令】：你浏览了博文和评论区，决定当一个“沉默的潜水者”。
            你的 actions 数组必须且只能是 ["view_only"]！
            请在 thought 中写下你真实的内心发散联想，但 content 必须为 null。"""
            rule_social = "3. **保持沉默**：你现在是潜水者，绝对不要发表任何回复或评论。"
        else:
            # 严格解耦两步意愿，强制执行
            post_req = '【强烈】=> 你必须选择 "like"(点赞原博)、"forward"(转发) 或 "comment"(直接评论原博，target_id填null) 至少一项！' if interact_post else '【无意愿】=> 绝不要对原博进行点赞、转发或直接评论！'
            comment_req = '【强烈】=> 你必须阅读前排评论，并选择 "like_comment"(点赞评论) 或 "comment"(回复他人，target_id填对方ID) 至少一项！' if interact_comment else '【无意愿】=> 绝对不要点赞任何评论，也绝对不要回复任何其他人！'

            action_prompt = f"""【行为指引】：群体倾向于【{bias_str}】。
            你当前的互动意愿已被系统严格拆分为“对原博”和“对评论区”两个独立部分，这两个意愿互不干扰，请【严格、独立】遵照执行：

            - 你对原博文的互动意愿：{post_req}
            - 你对评论区的盖楼意愿：{comment_req}

            可选的 actions 动作库（你可以自由组合数组，以同时满足上述两个意愿约束）：
            - "like" (点赞原博)
            - "like_comment" (点赞前排评论，需在liked_comment_ids填写目标ID)
            - "comment" (发表短评论或回复)
            - "forward" (纯转发扩散)
            - "forward_with_comment" (带评转发)"""

            if interact_comment:
                rule_social = "3. **积极盖楼（恢复网民社交）**：请务必仔细阅读【前排评论快照】！寻找共鸣，并在 content 中写下对他的回复（并在 target_id 中严格填入对方的纯数字ID），或者点赞他的评论。"
            else:
                rule_social = "3. **无视评论区**：你只关注原博，对别人的评论毫无兴趣，绝对不要点赞评论，也绝对不要回复别人（若要评论，必须是直接评论原博，且 target_id 必须填 null）。"

        if is_local_mode:
            system_prompt = f"""你现在扮演一名生活在【{location}】的真实的微博网民。
        【个人档案】：{json.dumps(local_persona, ensure_ascii=False)}

        {action_prompt}

        【🚨 拟真互动八大铁律】：
        1. **强制本地人视角**：无论你的档案原本写的是哪里，你现在的唯一居住地、工作地就是【{location}】！通报里的事就发生在你家门口！你必须从【{location}本地人】的切身利益出发进行思考和评价。
        2. **前台自然与短平快**：真实的公开评论非常简短！如果选择了评论，你的 content 必须严格控制在 10~40 个字以内！并且，您是【{location}】本地人，请以本地人的口吻评价、描述。
        {rule_social}
        4. **去演戏感**：评论中绝对禁止出现“作为一名XX（职业）”或“作为{location}人”的生硬自我介绍句式，请像普通路人一样自然表达。
        5. **温和防杠**：无论是回复别人还是评论原博，公开语气必须是正面支持原微博、中立吃瓜或善意期盼，绝不抬杠。
        6. **后台无限发散（真实OS解封）**：你的内心潜台词（thought）【不需要简短】！请深度发散通报对你这个本地居民的切身利益影响。
        7. **避免复读机**：整个评论区不能有大段大段相同评论刷屏。
        8. **限制讨论链与聚焦核心**：一个评论被讨论不能超过3条；所有评论都必须聚焦原微博，并且体现支持或者中立、温和建议的态度。
        必须严格输出 JSON 格式：
        {{
            "thought": "真实的潜意识(长篇大论，深度发散本地影响，不受字数限制)",
            "actions": ["动作1", "动作2"],
            "target_id": "想要回复的AgentID数字(若直接评论原博则填null)",
            "liked_comment_ids": ["想要点赞的评论的AgentID数字", "可填多个", "没有则留空数组"],
            "content": "公开评论或null(如果包含comment动作，此处限10~40字)",
            "trust_change": 0.1
        }}"""
        else:
            system_prompt = f"""你现在扮演一名的微博网民。
        【个人档案】：{json.dumps(local_persona, ensure_ascii=False)}

        {action_prompt}

        【🚨 拟真互动八大铁律】：
        1. **前台自然与短平快**：真实的公开评论非常简短！如果选择了评论，你的 content 必须严格控制在 10~40 个字以内！
        2. **独立意愿执行**：你的行动必须严格符合【行为指引】中对原博和对评论区的独立意愿约束。
        {rule_social}
        4. **去演戏感**：评论中绝对禁止出现“作为一名XX（职业）”的生硬句式，请像普通路人一样自然表达。
        5. **温和防杠**：无论是回复别人还是评论原博，公开语气必须是正面支持原微博、中立吃瓜或善意期盼，绝不抬杠。
        6. **后台无限发散（真实OS解封）**：你的内心潜台词（thought）【不需要简短】！可以发散、联想自身面临的焦虑。
        7. **避免复读机**：整个评论区不能有大段大段相同评论刷屏。
        8. **限制讨论链与聚焦核心**：一个评论被讨论不能超过 3 条；所有评论都必须聚焦原微博，并且体现支持或者中立、温和建议的态度。
        必须严格输出 JSON 格式：
        {{
            "thought": "真实的潜意识(长篇大论，深度发散，不受字数限制)",
            "actions": ["动作1", "动作2"],
            "target_id": "想要回复的AgentID数字(若直接评论原博则填null)",
            "liked_comment_ids": ["想要点赞的评论的AgentID数字", "可填多个", "没有则留空数组"],
            "content": "公开评论或null(如果包含comment动作，此处限10~40字)",
            "trust_change": 0.1
        }}"""

        try:
            res = self.client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user",
                     "content": f"【最新通报】：{post}\n\n{history}\n\n【前排评论快照(请积极寻找可回复的对象)】：\n{social_context}\n\n(提示：请结合你的性格，参考上述历史相似案例的【真实互动量】，来决定你这次的反应)"}
                ],
                response_format={"type": "json_object"},
                temperature=0.85
            )
            return json.loads(res.choices[0].message.content)
        except:
            return None


# ==========================================
# 5. 核心图表绘制引擎 (浅色背景版)
# ==========================================
def draw_dashboard_to_st(sim_data, time_labels, time_unit):
    set_matplotlib_font()
    fig = plt.figure(figsize=(16, 5.5), facecolor='white')
    gs = gridspec.GridSpec(1, 2, width_ratios=[2, 1], figure=fig)

    ax1 = fig.add_subplot(gs[0])
    ax1.set_facecolor('#F8F9FA')

    x_vals = [0] + time_labels
    y_exp = [0] + sim_data['exposure']
    y_int = [0] + sim_data['interaction']

    ax1.plot(x_vals, y_exp, 'o-', color='#00B4D8', linewidth=3, markersize=8, label='触达曝光总人数')
    ax1.plot(x_vals, y_int, 's-', color='#E63946', linewidth=3.5, markersize=8, label='实质互动(含评赞)')

    theoretical = [0] + [sim_data['num_agents'] * sim_data['prob'] * (1 - math.exp(-0.4 * i)) for i in range(1, 6)]
    ax1.plot(x_vals, theoretical, '--', color='#FFB703', alpha=0.9, linewidth=2.5, label='回归预测累积曲线')

    ax1.set_title(f"舆论场时间动力学 (T={sim_data['time_span']} {time_unit})", fontsize=16, fontweight='bold',
                  color='#333333', pad=15)
    ax1.set_xlabel(f"演化时间 ({time_unit})", fontsize=13, color='#333333')
    ax1.set_ylabel("受众数量 (人)", fontsize=13, color='#333333')
    ax1.set_xticks(x_vals)
    ax1.tick_params(labelsize=11, colors='#333333')
    ax1.legend(fontsize=12, facecolor='white', edgecolor='#CCCCCC', labelcolor='#333333')
    ax1.grid(color='#CCCCCC', linestyle=':', alpha=0.8)

    ax2 = fig.add_subplot(gs[1])
    ax2.set_facecolor('#F8F9FA')
    G = nx.DiGraph()
    G.add_edges_from(sim_data['edges'])
    if len(G.nodes) > 0:
        pos = nx.spring_layout(G, k=0.8, seed=42)
        nx.draw(G, pos, ax=ax2, node_size=60, node_color='#FFB703', edge_color='#999999', arrowsize=10, alpha=0.8)
    ax2.set_title("互动社交裂变图谱", fontsize=16, fontweight='bold', color='#333333', pad=15)

    plt.tight_layout()
    return fig


# ==========================================
# 6. 前端交互主视图
# ==========================================
st.title("🐟 WeiboFish：政务新媒体多智能体仿真沙盘")
st.markdown("基于 **32万条政务博文数据** 与多智能体模拟，快速推演政务新媒体的典型舆论情境")

prov_city_map, city_vol_map, prov_vol_map, global_vol = load_agenda_data()

with st.sidebar:
    st.header("⚙️ 基础引擎配置")
    api_key = st.text_input(
        "DeepSeek API Key",
        type="password",
        placeholder="请输入您的 sk-... 密钥",
        help="本系统不存储您的密钥，仅用于本次推演调用。您可以前往 DeepSeek 官网申请。"
    )

    total_calls = get_total_usage()
    st.sidebar.metric("📊 系统已累计推演", f"{total_calls} 次")
    is_local_mode = st.toggle("开启本地环境模式", value=True,
                              help="开启后网民将代入事发地居民视角，衰减按照默认设置；关闭则视为全网泛泛关注，衰减速率变缓。")
    num_agents = st.slider("注入网民智能体数量", min_value=10, max_value=500, value=100, step=10)
    time_span = st.slider("推演现实时间跨度", min_value=1, max_value=7, value=3, step=1)

    st.divider()
    st.header("🌍 地域与议程自动匹配")
    provinces = list(prov_city_map.keys())
    selected_prov = st.selectbox("1. 选择所属省份", provinces) if provinces else "默认省份"
    cities = prov_city_map.get(selected_prov, []) + ["其他 (使用省均值)"] if provinces else ["默认城市"]
    selected_city = st.selectbox("2. 选择地级市", cities)

    if selected_city == "其他 (使用省均值)":
        agenda_vol = prov_vol_map.get(selected_prov, global_vol)
        city_name = f"{selected_prov} (全省范围)"
    else:
        agenda_vol = city_vol_map.get(selected_city, global_vol)
        city_name = selected_city
    st.info(f"📊 **自动匹配议程波动值**: `{agenda_vol:.4f}`")

    st.divider()
    st.header("📱 媒体组合选项")
    media_options = st.multiselect("选择附件 (纯文本默认无需勾选)",
                                   ["图片 ", "视频", "网页链接", "话题、超话/At用户"])
    media_score = sum([1.5 if "视频" in m else 1.0 if "图片" in m else 0.5 for m in media_options])
    media_level = min(3.0, media_score)
    st.info(f"🧮 **计算得出媒体丰富度**: `{media_level:.1f}`")

post_content = st.text_area("✍️ 拟发布的政务微博内容", height=120, placeholder="在此输入政务通报正文...")

if st.button("🚀 运行 weibofish 实证推演", use_container_width=True, type="primary"):
    if not api_key:
        st.error("请先在左侧输入 API Key！")
        st.stop()
    if not post_content:
        st.warning("请输入微博通报正文！")
        st.stop()
    log_usage()
    client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")
    nlp, stats_model, mem, personas = load_ai_engines()
    context_text = f"【地域：{city_name}】{post_content}"

    # --- 阶段一：实证模型预测 ---
    with st.spinner("正在运行 RoBERTa 及多元回归模型..."):
        scores = nlp.analyze(context_text)
        z_read = (scores['readability_0_100'] - 65) / 12
        z_emo = (scores['emotion_0_100'] - 30) / 20
        perf = stats_model.calculate_excess_performance(z_read, z_emo, media_level, agenda_vol)

        if perf > 1.0:
            perf_eval_str = "极具爆款潜质，有望引发全网广泛关注"
        elif perf > 0:
            perf_eval_str = "有一定爆款潜质，表现将优于平均水平"
        elif perf > -1.0:
            perf_eval_str = "表现平平，属于常规政务通报水平"
        else:
            perf_eval_str = "不太可能成为爆款，预计关注度较低"

        act_prob = max(0.15, min(0.45, 0.15 + 0.30 / (1 + math.exp(-perf))))

        if z_emo > 0.5:
            bias_str = "情绪较高，倾向于评论和带评转发"
        elif z_read > 0.5:
            bias_str = "浅层阅读，倾向于纯点赞或纯转发"
        else:
            bias_str = "日常浏览，跟风点赞为主"

        # --- 阶段二：历史相似案卷检索 (RAG) ---
        st.markdown("---")
        st.subheader("一、历史相似案卷检索 (RAG)")

        # 1. 提前进行检索
        related = mem.retrieve_similar(context_text, top_k=3)

        # 2. 在独立的 UI 容器中展示结果 (不在分栏里了，独占一行)
        if related:
            st.success(f"🔍 记忆唤醒：成功从 32 万条历史语料中，匹配到 {len(related)} 条相似的真实案例！")
            with st.expander("🧠 点击展开查看详细历史卷宗", expanded=True):
                for i, c in enumerate(related):
                    st.markdown(f"**【案例 {i + 1}】** 匹配度: `{c['score']:.4f}`")
                    st.markdown(
                        f"👤 **发布账号**：{c.get('account', '未知')} &nbsp;&nbsp;|&nbsp;&nbsp; 🕒 **时间**：{c.get('date', '未知')}")
                    st.markdown(f"📊 **真实互动**：`{c.get('engagement', '无数据')}`")
                    # 嵌套一个默认折叠的面板来展示完整原文
                    with st.expander("📝 点击查看完整原文", expanded=False):
                        st.info(c.get('content', '无内容'))
                    if i < len(related) - 1:
                        st.divider()
        else:
            st.info("💡 系统中暂未匹配到极度相似的历史案例，Agent 将依靠基础社会常识进行推演。")

        # 3. 把记忆组装成字符串，准备喂给下面的大模型
        if related:
            history_str = "【历史相似案例与真实反响参考】\n" + "\n".join([
                f"-> 时间：{c.get('date', '未知')} | 账号：【{c.get('account', '未知')}】\n"
                f"-> 真实互动量：{c.get('engagement', '无数据')}\n"
                f"-> 相似内容：{c.get('content', '无内容')[:60]}...\n"
                for c in related
            ])
        else:
            history_str = "【历史相似案例参考】：暂无极度相似的历史通报。"

        # --- 阶段三：多智能体推演实况 (MAS) ---
        st.markdown("---")
        st.subheader("二、微观群体行为漏斗与潜意识透视")

        col_chat, col_os = st.columns([1.2, 1.2])
        with col_chat:
            st.caption("🗣️ **前台：显性互动区** (点赞、转发与盖楼回复)")
            chat_box = st.container(height=550)
        with col_os:
            st.caption("💭 **后台：潜意识监控区** (无拘无束的真实心理活动)")
            os_box = st.container(height=550)
    # =================按比例抽样逻辑开始=================
    stance_pools = {}
    for p in personas:
        stype = p.get('psychology', {}).get('stance_type', '其他')
        if stype not in stance_pools:
            stance_pools[stype] = []
        stance_pools[stype].append(p)

    sampled_personas = []
    total_personas = len(personas)
    remaining_spots = num_agents
    for stype, pool in stance_pools.items():
        quota = int((len(pool) / total_personas) * num_agents)
        if quota > 0:
            sampled_personas.extend(random.choices(pool, k=quota) if quota > len(pool) else random.sample(pool, quota))
        remaining_spots -= quota

    if remaining_spots > 0:
        sampled_personas.extend(random.choices(personas, k=remaining_spots))

    random.shuffle(sampled_personas)
    agents = [StreamlitAgent(i, sampled_personas[i], client) for i in range(num_agents)]
    # =================按比例抽样逻辑结束=================

    time_unit = "小时" if time_span == 1 else "天"
    max_time = 24 if time_span == 1 else time_span
    time_labels = [round((i + 1) / 5 * max_time, 1) for i in range(5)]

    sim_data = {'steps': [], 'exposure': [], 'interaction': [], 'edges': [], 'num_agents': num_agents, 'prob': act_prob,
                'time_span': max_time}
    comments_pool, full_logs, thoughts_pool = [], [], []
    stats = {"view_only": 0, "like": 0, "comment": 0, "forward": 0, "forward_c": 0, "like_comment": 0}

    comments_data = {}

    decay_constant = 0.35 if is_local_mode else 0.28

    progress = st.progress(0)
    for t in range(5):
        current_time_str = f"{time_labels[t]} {time_unit}"
        progress.progress((t + 1) * 20, text=f"舆论推演时间线：第 {current_time_str} / {max_time} {time_unit}...")

        decay_factor = math.exp(-decay_constant * t)

        base_post_prob = act_prob * decay_factor
        base_comment_prob = min(0.6, base_post_prob * 1.5)

        if t < 3:
            algo_push = int(num_agents * 0.25 * decay_factor)
            unexposed = [a for a in agents if not a.is_exposed]
            for a in random.sample(unexposed, min(algo_push, len(unexposed))): a.is_exposed = True

        current_active = [a for a in agents if a.is_exposed and not a.has_interacted]
        os_count_this_wave = 0

        for a in current_active:
            interact_post = random.random() < base_post_prob
            interact_comment = (random.random() < base_comment_prob) if len(comments_pool) > 0 else False

            is_interacting = interact_post or interact_comment

            if not is_interacting:
                stats['view_only'] += 1
                if os_count_this_wave < max(3, num_agents // 50):
                    os_count_this_wave += 1
                else:
                    a.has_interacted = True
                    continue

            social = " | ".join(comments_pool[-8:])
            res = a.react(context_text, history_str, social, bias_str, interact_post, interact_comment, city_name,
                          is_local_mode)

            if res and res.get('actions'):
                acts = res['actions'] if isinstance(res['actions'], list) else [res['actions']]
                thought = res.get('thought', '')
                target = extract_id(res.get('target_id'))
                target_str = f" 回复 @Agent_{target:02d}" if target is not None else ""
                content = res.get('content', '')
                has_real_action = False

                liked_ids = res.get('liked_comment_ids', [])
                if isinstance(liked_ids, list):
                    for lid in liked_ids:
                        lid_int = extract_id(str(lid))
                        if lid_int is not None and lid_int in comments_data:
                            comments_data[lid_int]['likes'] += 1
                            stats['like_comment'] += 1
                            with chat_box: st.info(
                                f"👍 **{a.role}** (Agent_{a.agent_id:02d}) 赞了 Agent_{lid_int:02d} 的评论。")
                            has_real_action = True

                if thought:
                    thoughts_pool.append(f"[{a.role} Agent_{a.agent_id}]: {thought}")
                    os_style = "os-view" if not is_interacting else ""
                    with os_box:
                        st.markdown(
                            f"<div class='os-box {os_style}'>🧠 <b>Agent_{a.agent_id:02d} ({a.role})</b><br>OS: {thought}</div>",
                            unsafe_allow_html=True)

                if is_interacting:
                    if 'like' in acts:
                        stats['like'] += 1
                        with chat_box: st.info(f"👍 **{a.role}** (Agent_{a.agent_id:02d}) 赞了该微博。")
                        has_real_action = True

                    if 'forward' in acts or 'forward_with_comment' in acts:
                        stats['forward'] += 1
                        if 'forward_with_comment' in acts and content:
                            stats['forward_c'] += 1
                            log_str = f"[{a.role} Agent_{a.agent_id:02d}]{target_str}: {content}"
                            with chat_box:
                                st.warning(
                                    f"🔁🔥 **{a.role}** (Agent_{a.agent_id:02d}) 带评转发{target_str}:\n\n“{content}”")
                            comments_pool.append(f"Agent{a.agent_id}: {content}")
                            full_logs.append(log_str)

                            comments_data[a.agent_id] = {"role": a.role, "traits": a.persona.get('psychology', {}).get(
                                'personality_traits', '普通网民'), "content": content, "likes": 0, "target": target}

                            if target is not None and target < len(agents): sim_data['edges'].append(
                                (a.agent_id, target))
                        else:
                            with chat_box:
                                st.success(f"🔁 **{a.role}** (Agent_{a.agent_id:02d}) 转发扩散了该内容。")
                        has_real_action = True

                    if 'comment' in acts and content and 'forward_with_comment' not in acts:
                        stats['comment'] += 1
                        log_str = f"[{a.role} Agent_{a.agent_id:02d}]{target_str}: {content}"
                        with chat_box:
                            st.chat_message("user", avatar="💬").write(
                                f"**{a.role}** (Agent_{a.agent_id:02d}){target_str}: \n\n {content}")
                        comments_pool.append(f"Agent{a.agent_id}: {content}")
                        full_logs.append(log_str)

                        comments_data[a.agent_id] = {"role": a.role,
                                                     "traits": a.persona.get('psychology', {}).get('personality_traits',
                                                                                                   '普通网民'),
                                                     "content": content, "likes": 0, "target": target}

                        if target is not None and target < len(agents): sim_data['edges'].append((a.agent_id, target))
                        has_real_action = True

                    if has_real_action:
                        current_exposure = sum(1 for x in agents if x.is_exposed)
                        fission_max = max(1, int((current_exposure / 30) * decay_factor))
                        for _ in range(random.randint(0, fission_max)):
                            f = random.choice(agents)
                            f.is_exposed = True
                            sim_data['edges'].append((a.agent_id, f.agent_id))
            a.has_interacted = True

        sim_data['steps'].append(t)
        sim_data['exposure'].append(sum(1 for a in agents if a.is_exposed))
        sim_data['interaction'].append(sum(1 for a in agents if a.has_interacted) - stats['view_only'])

    # --- 阶段三：核心数据看板 ---
    st.markdown("---")
    st.subheader(f"三、舆论场漏斗数据看板 (追踪时效：{max_time} {time_unit})")

    col1, col2, col3, col4, col5 = st.columns(5)
    col1.metric("👁️ 样本曝光人数", f"{sim_data['exposure'][-1]} 人")
    col2.metric("👤 仅浏览(潜水)", f"{stats['view_only']} 人")
    col3.metric("👍 样本赞博/赞评", f"{stats['like']} / {stats['like_comment']} 次")
    col4.metric("🔁 样本转发", f"{stats['forward']} 次")
    col5.metric("💬 样本评论", f"{stats['comment'] + stats['forward_c']} 条")

    st.pyplot(draw_dashboard_to_st(sim_data, time_labels, time_unit), use_container_width=True)

    # --- 阶段 3.5：模拟互动评论区还原 ---
    st.markdown("---")
    st.subheader("💬 四、模拟博文及评论区互动还原")
    st.caption("真实呈现前台盖楼与点赞情况（含参与网民之隐性特征标注）：")

    html_str = ""
    root_cids = [cid for cid, cinfo in comments_data.items() if
                 cinfo['target'] is None or cinfo['target'] not in comments_data]

    for root_id in root_cids:
        root_info = comments_data[root_id]
        html_str += f"""<div style="background-color: #ffffff; padding: 15px; border-radius: 8px; margin-bottom: 15px; border: 1px solid #e0e0e0;">
<div style="font-weight: bold; color: #eb7350; font-size: 14px; margin-bottom: 5px;">{root_info['role']} (Agent_{root_id:02d}) <span style="font-weight: normal; color: #999; font-size: 12px;">[{root_info['traits']}]</span></div>
<div style="color: #333; font-size: 15px; line-height: 1.5;">{root_info['content']}</div>
<div style="color: #808080; font-size: 13px; margin-top: 8px;">👍 {root_info['likes']} &nbsp;&nbsp;💬 回复</div>
"""
        descendants = []


        def get_descendants(parent_id):
            children = [cid for cid, cinfo in comments_data.items() if cinfo['target'] == parent_id]
            for child_id in children:
                descendants.append(child_id)
                get_descendants(child_id)


        get_descendants(root_id)

        for r_id in descendants:
            r_info = comments_data[r_id]
            target_id = r_info['target']
            prefix = ""
            if target_id != root_id and target_id in comments_data:
                target_role = comments_data[target_id]['role']
                prefix = f"<span style='color: #1c4e7d; margin-right: 5px;'>回复 @{target_role} (Agent_{target_id:02d}):</span>"

            html_str += f"""<div style="background-color: #f9f9f9; padding: 12px; margin-top: 10px; margin-left: 20px; border-radius: 6px; border-left: 3px solid #00b4d8;">
<div style="font-weight: bold; color: #eb7350; font-size: 14px; margin-bottom: 5px;">{r_info['role']} (Agent_{r_id:02d}) <span style="font-weight: normal; color: #999; font-size: 12px;">[{r_info['traits']}]</span></div>
<div style="color: #333; font-size: 15px; line-height: 1.5;">{prefix}{r_info['content']}</div>
<div style="color: #808080; font-size: 13px; margin-top: 8px;">👍 {r_info['likes']}</div>
</div>
"""
        html_str += "</div>\n"

    if not html_str:
        html_str = "<div style='color:#999; text-align:center; padding:20px;'>暂无评论</div>"

    post_html = f"""<div style="background-color: #ffffff; padding: 20px; border-radius: 10px; border: 1px solid #e0e0e0; box-shadow: 0 2px 5px rgba(0,0,0,0.05); margin-bottom: 10px;">
<div style="display: flex; align-items: center; margin-bottom: 15px;">
<div style="width: 45px; height: 45px; background-color: #1c4e7d; border-radius: 50%; display: flex; justify-content: center; align-items: center; color: white; font-weight: bold; font-size: 18px; margin-right: 15px;">政</div>
<div>
<div style="font-weight: bold; color: #333; font-size: 16px;">{city_name}某政务微博账户 <span style="color: #1DA1F2; font-size: 14px;">✔️蓝V认证</span></div>
<div style="color: #999; font-size: 12px;">刚刚 发布</div>
</div>
</div>
<div style="color: #333; font-size: 16px; line-height: 1.8; margin-bottom: 15px;">
{post_content.replace(chr(10), '<br>')}
</div>
<div style="color: #888; font-size: 14px; display: flex; gap: 30px; border-top: 1px solid #f0f0f0; padding-top: 15px; padding-bottom: 15px;">
<span>🔁 转发 {stats['forward']}</span>
<span>💬 评论 {stats['comment'] + stats['forward_c']}</span>
<span>👍 点赞 {stats['like']}</span>
</div>
<details style="border-top: 1px dashed #e0e0e0; padding-top: 15px; cursor: pointer;">
<summary style="color: #1c4e7d; font-weight: bold; outline: none; margin-bottom: 15px;">点击展开评论区</summary>
{html_str}
</details>
</div>
"""
    st.markdown(post_html, unsafe_allow_html=True)

    # --- 阶段四：政务智库研判专报 ---
    st.markdown("---")
    st.subheader("五、政务内参：舆论场心理诊断")
    with st.spinner("智库模型正在对齐实证预测与沙盘结果，深度生成应对策略..."):
        prompt = f"""你是一名专门为政府高层提供核心内参的顶级社会学与数据分析专家。
                【原始通报文本】：{post_content}
                【事件发生地】：{city_name}

                [理论定量数据]：
                【论文实证预估互动率基准】：{act_prob:.1%}
                【模型理论预测结论】：该博文{perf_eval_str}
                【追踪时效】：{max_time} {time_unit}

                [历史记忆与对标案卷]：
                {history_str}

                [沙盘抽样演化数据]（注：本次推演仅为有限智能体抽样所展现的可能情境之一）：
                【抽样曝光总人数】：{sim_data['exposure'][-1]}
                【实际互动总人数】：{sim_data['interaction'][-1]} (原博点赞{stats['like']}, 评论点赞{stats['like_comment']}, 转发{stats['forward']}, 评论{stats['comment']})
                【仅浏览不互动的潜水者】：{stats['view_only']}

                [定性语料数据]：
                【前台-公开互动与盖楼记录】：{full_logs}
                【后台-网民潜意识OS】：{thoughts_pool}

                【排版红线要求】：
                1. 绝对禁止使用任何 Markdown 星号（*）！
                2. 请使用全角中文标点，使用“一、二、三、”作为主标题，“1. 2. 3.”作为子标题，保持排版的干净、严肃（注意：数字、百分号及数字中的小数点必须保持半角，如“35.2%”，严禁写成“35。2％”）。

                【核心诊断任务】：
                1. **历史案卷对标与经验萃取**：严格对照提供的【历史记忆与对标案卷】（如果有的话），分析历史上同类通报的真实网民互动表现。指出历史经验对本次事件处置的借鉴意义（例如历史通报是成功平息了舆论，还是引发了次生灾害）。
                2. **潜质定调与偏差解释**：结合【模型理论预测结论】与本次【沙盘抽样演化数据】进行对比。指出由于现实中微博环境复杂，当前沙盘展现的“典型情境”与理论预估之间存在何种偏差及原因。
                3. **本地群体心态与高赞评论研判**：分析{city_name}当地网民的前台支持与后台OS的温差，特别注意【评论点赞数】和盖楼互动情况，指出潜在的社会风险点。
                4. **针对性文本重构策略**：基于以上所有分析（特别是历史翻车或成功的经验），给出非常具体的通报文本修改建议。
                """

        obs = client.chat.completions.create(
            model="deepseek-reasoner",
            messages=[{"role": "user", "content": prompt}]
        )

        clean_report = obs.choices[0].message.content.replace('*', '')

        st.markdown(f"""
        <div class="report-box">
            <div class="report-title">【决策内参】针对{city_name}某政务微博账户的舆论演化动力学专报</div>
            {clean_report}
        </div>
        """, unsafe_allow_html=True)

        with st.expander("🧐 展开查看：内参推理的 Chain of Thought"):
            st.write(obs.choices[0].message.reasoning_content)

        try:
            word_bytes = create_word_report(city_name, clean_report)
            import base64

            b64 = base64.b64encode(word_bytes).decode()
            href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{city_name}舆情内参专报.docx" style="display: inline-block; padding: 0.6em 1.2em; color: white; background-color: #ff4b4b; border-radius: 4px; text-decoration: none; font-weight: bold; margin-top: 10px;">📥 下载专报 (Word格式)</a>'
            st.markdown(href, unsafe_allow_html=True)
        except Exception as e:
            st.error(f"Word导出按钮加载失败: {e}")

        st.session_state.sim_completed = True
        
# ==========================================
# 7. 意见与反馈模块 (推演完成后才会显示)
# ==========================================
if st.session_state.sim_completed:
    st.markdown("---")
    st.subheader("📝 意见与反馈")
    st.write("您的真实体验对本政务新媒体推演平台的迭代非常重要！")

    with st.form("feedback_form", clear_on_submit=True):
        # 1. 姓名/称呼（新增：必填项）
        user_name = st.text_input("您的称呼/姓名（必填）：", placeholder="请输入您的名字或昵称")
        
        # 2. 星星打分（必选项）
        star_mapping = {"⭐ (1分)": 1, "⭐⭐ (2分)": 2, "⭐⭐⭐ (3分)": 3, "⭐⭐⭐⭐ (4分)": 4, "⭐⭐⭐⭐⭐ (5分)": 5}
        selected_stars = st.radio(
            "您对本次多智能体推演的整体满意度（必选）：", 
            options=list(star_mapping.keys()), 
            index=None, 
            horizontal=True
        )
        
        # 3. 文字反馈（选填）
        feedback_text = st.text_area("请详细描述您的建议或遇到的问题（选填）：", placeholder="例如：智能体生成的应对策略是否符合实际情况？")
        
        # 4. 邮箱（选填）
        user_email = st.text_input("您的联系邮箱（选填）：", placeholder="方便我们后续与您交流探讨")
        
        submitted = st.form_submit_button("发送反馈")

        if submitted:
            if not user_name.strip():
                st.warning("提交失败：请填写您的称呼/姓名后再提交。")
            elif selected_stars is None:
                st.warning("提交失败：请先给本次推演打个分（点击上方星星）哦！")
            else:
                import os
                import pandas as pd
                from datetime import datetime
                
                rating = star_mapping[selected_stars] 
                time_now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

                safe_feedback = feedback_text.strip() if feedback_text.strip() else "未填写"
                safe_email = user_email.strip() if user_email.strip() else "未填写"

                new_data = pd.DataFrame({
                    "时间": [time_now], 
                    "反馈者": [user_name.strip()], # 新增列
                    "满意度评分": [rating], 
                    "反馈内容": [safe_feedback],
                    "联系邮箱": [safe_email] 
                })
                
                feedback_file = "feedback.csv"
                if not os.path.exists(feedback_file):
                    new_data.to_csv(feedback_file, index=False, encoding='utf-8-sig') 
                else:
                    new_data.to_csv(feedback_file, mode='a', header=False, index=False, encoding='utf-8-sig')
                    
                st.success(f"🎉 感谢您的反馈，{user_name.strip()}！我已经拿小本本记下了。")